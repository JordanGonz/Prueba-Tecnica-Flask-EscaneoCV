import logging
from docx import Document

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path):
    """
    Extract text content from a DOCX file.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text content
        
    Raises:
        Exception: If DOCX processing fails
    """
    try:
        extracted_text = ""
        
        # Load the document
        doc = Document(file_path)
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                extracted_text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    extracted_text += " | ".join(row_text) + "\n"
        
        # Extract text from headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        extracted_text += paragraph.text + "\n"
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        extracted_text += paragraph.text + "\n"
        
        if not extracted_text.strip():
            raise Exception("No text could be extracted from the DOCX file")
            
        logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX: {file_path}")
        return extracted_text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        raise Exception(f"Failed to process DOCX file: {str(e)}")

def extract_text_from_docx_bytes(docx_bytes):
    """
    Extract text content from DOCX bytes.
    
    Args:
        docx_bytes (bytes): DOCX file content as bytes
        
    Returns:
        str: Extracted text content
        
    Raises:
        Exception: If DOCX processing fails
    """
    try:
        from io import BytesIO
        
        extracted_text = ""
        
        # Load the document from bytes
        docx_stream = BytesIO(docx_bytes)
        doc = Document(docx_stream)
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                extracted_text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    extracted_text += " | ".join(row_text) + "\n"
        
        if not extracted_text.strip():
            raise Exception("No text could be extracted from the DOCX file")
            
        logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX bytes")
        return extracted_text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX bytes: {str(e)}")
        raise Exception(f"Failed to process DOCX file: {str(e)}")
